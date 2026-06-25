#!/usr/bin/env python3
"""FinClaw-parse-upload

Parse user-uploaded financial statement files into a uniform normalized JSON
that the rest of the FinClaw pipeline (foundation / business analysis) can
consume.

Supports (with explicit support levels, never silently):
  - supported now:
      * .xlsx, .xlsm      (openpyxl)
      * .csv, .tsv        (csv stdlib)
      * .pdf   (text-based; pymupdf)         -- best-effort
      * .docx             (python-docx)       -- best-effort
      * .zip              (zipfile, recursive into any of the above)
  - supported via external skill (when available, marker-pdf / ocr-and-documents):
      * scanned PDF, image scans (.png / .jpg / .jpeg)
      * legacy .xls (xlrd needed)
  - rejected explicitly (never fabricated):
      * anything outside the above, or with binary content we cannot decode

Output schema (printed to stdout as JSON):
{
  "files": [
    {
      "input_path": "...",
      "kind": "xlsx|csv|tsv|pdf|docx|zip|rejected",
      "status": "ok|partial|rejected",
      "support_level": "core|optional|delegated|rejected",
      "needs_dependencies": [],
      "warnings": [],
      "error": null,
      "pages_or_sheets": 0,
      "sheets": [ { "name": "...", "rows": N, "cols": M, "preview": [[...]] } ],
      "text_excerpt": "...",          # for pdf/docx/zip
      "extracted_files": [ ... ]      # for zip
    }
  ],
  "summary": { "ok": n, "partial": n, "rejected": n }
}

Usage:
  python3 FinClaw-parse-upload.py <file_or_dir> [<file_or_dir> ...]
  python3 FinClaw-parse-upload.py --json-out out.json <inputs...>
"""
from __future__ import annotations
import csv
import io
import json
import os
import sys
import zipfile
import shutil
import tempfile
import traceback
from pathlib import Path
from typing import Any

SUPPORTED_EXTS = {
    ".xlsx": "core",
    ".xlsm": "core",
    ".csv": "core",
    ".tsv": "core",
    ".pdf": "optional",
    ".docx": "optional",
    ".zip": "core",
    ".png": "delegated",
    ".jpg": "delegated",
    ".jpeg": "delegated",
    ".xls": "delegated",
}

PREVIEW_ROWS = 5
PREVIEW_COLS = 12
TEXT_EXCERPT_CHARS = 4000

# ---- dependency probes -----------------------------------------------------

def has_mod(name: str) -> bool:
    try:
        __import__(name)
        return True
    except Exception:
        return False

# ---- per-kind parsers -------------------------------------------------------

def _truncate_rows(rows):
    out = []
    for r in rows[:PREVIEW_ROWS]:
        # cast each cell to str, truncate length
        row = []
        for c in r[:PREVIEW_COLS]:
            if c is None:
                row.append("")
            else:
                s = str(c)
                if len(s) > 200:
                    s = s[:200] + "…"
                row.append(s)
        out.append(row)
    return out

def parse_xlsx(path: Path) -> dict:
    if not has_mod("openpyxl"):
        return {
            "status": "rejected",
            "support_level": "core",
            "needs_dependencies": ["openpyxl"],
            "error": "缺少依赖 openpyxl，无法解析 xlsx/xlsm。请联系管理员安装。",
        }
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    sheets = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append(list(row))
        sheets.append({
            "name": name,
            "rows": len(rows),
            "cols": ws.max_column or 0,
            "preview": _truncate_rows(rows),
        })
    wb.close()
    return {
        "status": "ok",
        "support_level": "core",
        "pages_or_sheets": len(sheets),
        "sheets": sheets,
    }

def parse_csv_tsv(path: Path, ext: str) -> dict:
    delimiter = "\t" if ext == ".tsv" else ","
    # encoding sniff
    text = None
    used_enc = None
    for enc in ("utf-8-sig", "utf-8", "gbk", "gb18030"):
        try:
            text = path.read_text(encoding=enc)
            used_enc = enc
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        return {
            "status": "rejected",
            "support_level": "core",
            "error": "文件编码无法识别（尝试了 utf-8/gbk）。请另存为 UTF-8 或 GBK 后重试。",
        }
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = [r for r in reader]
    sheets = [{
        "name": path.stem,
        "rows": len(rows),
        "cols": max((len(r) for r in rows), default=0),
        "preview": _truncate_rows(rows),
        "encoding": used_enc,
    }]
    return {
        "status": "ok",
        "support_level": "core",
        "pages_or_sheets": 1,
        "sheets": sheets,
    }

def parse_pdf(path: Path) -> dict:
    if not has_mod("pymupdf"):
        return {
            "status": "rejected",
            "support_level": "optional",
            "needs_dependencies": ["pymupdf"],
            "error": "缺少依赖 pymupdf，无法解析 PDF。请联系管理员安装，或把 PDF 转为 Excel/CSV 再上传。",
        }
    import pymupdf
    doc = pymupdf.open(path)
    pages_text = []
    for p in doc:
        pages_text.append(p.get_text() or "")
    doc.close()
    full_text = "\n\n".join(pages_text).strip()
    if not full_text:
        return {
            "status": "partial",
            "support_level": "delegated",
            "pages_or_sheets": len(pages_text),
            "needs_dependencies": ["marker-pdf", "rapidocr-onnxruntime", "pytesseract+tesseract"],
            "warnings": ["PDF 未检测到文本层，可能是扫描件。需要 OCR 才能提取数字。"],
            "error": "此 PDF 为扫描件/无文本层，需要 OCR 依赖（marker-pdf ~5GB / rapidocr / tesseract）才能解析。请先在系统安装这些依赖，或将扫描件先转成 Excel/CSV。",
        }
    excerpt = full_text[:TEXT_EXCERPT_CHARS]
    return {
        "status": "ok",
        "support_level": "optional",
        "pages_or_sheets": len(pages_text),
        "sheets": [],
        "text_excerpt": excerpt,
        "warnings": [] if len(full_text) <= TEXT_EXCERPT_CHARS else [f"PDF 文本较长（{len(full_text)} 字），仅展示前 {TEXT_EXCERPT_CHARS} 字"],
    }

def parse_docx(path: Path) -> dict:
    if not has_mod("docx"):
        return {
            "status": "rejected",
            "support_level": "optional",
            "needs_dependencies": ["python-docx"],
            "error": "缺少依赖 python-docx，无法解析 Word。请联系管理员安装，或把 Word 转成 PDF/Excel 再上传。",
        }
    import docx
    d = docx.Document(str(path))
    paragraphs = [p.text for p in d.paragraphs if p.text]
    tables = []
    for ti, t in enumerate(d.tables):
        rows = [[c.text for c in row.cells] for row in t.rows]
        tables.append({"name": f"Table{ti+1}", "rows": len(rows), "cols": max((len(r) for r in rows), default=0), "preview": _truncate_rows(rows)})
    text = "\n".join(paragraphs)
    return {
        "status": "ok",
        "support_level": "optional",
        "pages_or_sheets": len(d.paragraphs),
        "sheets": tables,
        "text_excerpt": text[:TEXT_EXCERPT_CHARS],
    }

def parse_zip(path: Path) -> dict:
    extracted = []
    rejected_children = []
    try:
        with zipfile.ZipFile(path) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                inner_name = Path(info.filename).name
                inner_ext = Path(info.filename).suffix.lower()
                if inner_ext in SUPPORTED_EXTS and inner_ext != ".zip":
                    with tempfile.TemporaryDirectory() as td:
                        target = Path(td) / inner_name
                        with zf.open(info) as src, open(target, "wb") as dst:
                            shutil.copyfileobj(src, dst)
                        child = parse_one(target)
                        child["input_path"] = f"{path}::{info.filename}"
                        extracted.append(child)
                else:
                    rejected_children.append({
                        "filename": info.filename,
                        "reason": "压缩包内含未支持的文件类型" if inner_ext else "无后缀名",
                    })
        status = "ok"
        if any(c["status"] == "rejected" for c in extracted) or rejected_children:
            status = "partial"
        return {
            "status": status,
            "support_level": "core",
            "extracted_files": extracted,
            "warnings": ([f"跳过了 {len(rejected_children)} 个未支持类型"] if rejected_children else []),
            "rejected_children": rejected_children,
        }
    except zipfile.BadZipFile:
        return {
            "status": "rejected",
            "support_level": "core",
            "error": "文件不是合法的 ZIP 压缩包。",
        }

def parse_image(path: Path) -> dict:
    return {
        "status": "rejected",
        "support_level": "delegated",
        "needs_dependencies": ["marker-pdf", "rapidocr-onnxruntime", "pytesseract+tesseract"],
        "error": "图片扫描件需要 OCR 依赖（marker-pdf ~5GB / rapidocr / tesseract）才能解析。请先在系统安装这些依赖，或把图片先转成 Excel/CSV。",
    }

def parse_xls(path: Path) -> dict:
    if has_mod("xlrd"):
        try:
            import xlrd
            book = xlrd.open_workbook(str(path))
            sheets = []
            for s in book.sheets():
                rows = []
                for r in range(s.nrows):
                    rows.append([s.cell_value(r, c) for c in range(s.ncols)])
                sheets.append({"name": s.name, "rows": len(rows), "cols": s.ncols, "preview": _truncate_rows(rows)})
            return {
                "status": "ok",
                "support_level": "delegated",
                "pages_or_sheets": len(sheets),
                "sheets": sheets,
            }
        except Exception as e:
            return {
                "status": "rejected",
                "support_level": "delegated",
                "needs_dependencies": ["xlrd"],
                "error": f"xlrd 解析 .xls 失败：{e}",
            }
    return {
        "status": "rejected",
        "support_level": "delegated",
        "needs_dependencies": ["xlrd"],
        "error": "缺少依赖 xlrd，无法解析 .xls。请把文件另存为 .xlsx 再上传，或联系管理员安装 xlrd。",
    }

# ---- dispatcher -------------------------------------------------------------

def parse_one(path: Path) -> dict:
    if not path.exists():
        return {
            "input_path": str(path),
            "kind": "rejected",
            "status": "rejected",
            "support_level": "rejected",
            "error": f"文件不存在：{path}",
        }
    ext = path.suffix.lower()
    base = {
        "input_path": str(path),
        "kind": ext.lstrip(".") or "unknown",
        "support_level": SUPPORTED_EXTS.get(ext, "rejected"),
        "needs_dependencies": [],
        "warnings": [],
        "error": None,
        "pages_or_sheets": 0,
        "sheets": [],
        "text_excerpt": "",
        "extracted_files": [],
    }
    if ext not in SUPPORTED_EXTS:
        base.update({
            "kind": "rejected",
            "status": "rejected",
            "error": f"不支持的文件类型：{ext or '(无后缀)'}。已支持的类型：xlsx/xlsm, csv/tsv, pdf, docx, zip, png/jpg（OCR 依赖可选），xls（xlrd 依赖可选）。",
        })
        return base
    try:
        if ext in (".xlsx", ".xlsm"):
            res = parse_xlsx(path)
        elif ext in (".csv", ".tsv"):
            res = parse_csv_tsv(path, ext)
        elif ext == ".pdf":
            res = parse_pdf(path)
        elif ext == ".docx":
            res = parse_docx(path)
        elif ext == ".zip":
            res = parse_zip(path)
        elif ext in (".png", ".jpg", ".jpeg"):
            res = parse_image(path)
        elif ext == ".xls":
            res = parse_xls(path)
        else:
            res = {"status": "rejected", "error": f"未实现：{ext}"}
    except Exception as e:
        res = {
            "status": "rejected",
            "error": f"解析异常：{e}",
            "trace": traceback.format_exc(limit=3),
        }
    base.update(res)
    return base

def main(argv):
    args = [a for a in argv[1:] if not a.startswith("--")]
    flags = {a for a in argv[1:] if a.startswith("--")}
    if not args or "--help" in flags or "-h" in flags:
        print(__doc__)
        return 0
    json_out = None
    if "--json-out" in flags:
        i = argv.index("--json-out")
        json_out = argv[i + 1]

    inputs = []
    for a in args:
        p = Path(a)
        if p.is_dir():
            for sub in sorted(p.rglob("*")):
                if sub.is_file():
                    inputs.append(sub)
        else:
            inputs.append(p)

    results = [parse_one(p) for p in inputs]
    summary = {
        "ok": sum(1 for r in results if r["status"] == "ok"),
        "partial": sum(1 for r in results if r["status"] == "partial"),
        "rejected": sum(1 for r in results if r["status"] == "rejected"),
    }
    out = {"files": results, "summary": summary}
    text = json.dumps(out, ensure_ascii=False, indent=2)
    if json_out:
        Path(json_out).write_text(text, encoding="utf-8")
    print(text)
    # nonzero exit if anything rejected, so callers can react
    return 0 if summary["rejected"] == 0 else 3

if __name__ == "__main__":
    sys.exit(main(sys.argv))
